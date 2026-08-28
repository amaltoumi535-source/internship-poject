"""DOCX file text extraction using python-docx."""

from docx import Document
from pathlib import Path
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)

class DOCXParser:
    """Extract text from Word (.docx) documents."""

    @staticmethod
    def extract(file_path: Path) -> Dict[str, Any]:
        """
        Extract text from DOCX file.

        Extracts:
        - All paragraphs
        - Table content
        - Headings with structure

        Returns:
            {
                "text": "full extracted text with structure",
                "metadata": {
                    "total_paragraphs": int,
                    "total_tables": int,
                    "language": "multi",
                    "confidence_score": 0.99
                },
                "structured_content": {
                    "headings": [...],
                    "paragraphs": [...],
                    "tables": [...]
                }
            }
        """
        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs: list = []
            headings: list = []

            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    paragraphs.append(para.text)

                    # Detect headings (style starts with "Heading")
                    try:
                        style_name = para.style.name or ""
                    except Exception:
                        style_name = ""
                    if style_name.startswith("Heading"):
                        headings.append({
                            "level": style_name,
                            "text": para.text
                        })

            # Extract tables
            tables_data: list = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_cells)
                tables_data.append(table_rows)

            # Format full text
            full_text = "\n\n".join(paragraphs)

            return {
                "text": full_text,
                "metadata": {
                    "total_paragraphs": len(paragraphs),
                    "total_tables": len(tables_data),
                    "language": "multi",
                    "confidence_score": 0.99,  # Native DOCX is very reliable
                    "parser": "python-docx"
                },
                "structured_content": {
                    "headings": headings,
                    "paragraphs": paragraphs,
                    "tables": tables_data
                }
            }

        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise Exception(f"DOCX parsing failed: {str(e)}")