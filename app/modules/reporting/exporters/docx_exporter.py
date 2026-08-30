# app/modules/reporting/exporters/docx_exporter.py
"""DOCX format exporter."""

from typing import Any, Dict
from io import BytesIO

from .base import BaseExporter


class DOCXExporter(BaseExporter):
    """Export reports to DOCX (Word) format."""

    def export(self, data: Dict[str, Any], filename: str) -> bytes:
        """Export data as DOCX."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        doc = Document()
        
        # Title
        if "title" in data:
            title = doc.add_heading(data["title"], level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if "summary" in data:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(data["summary"])
            doc.add_paragraph()  # Blank line
        
        # Metadata
        if "metadata" in data:
            doc.add_heading("Metadata", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Key"
            hdr_cells[1].text = "Value"
            for key, value in data["metadata"].items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
            doc.add_paragraph()  # Blank line
        
        # Content/Results
        if "results" in data:
            doc.add_heading("Results", level=2)
            results = data["results"]
            if isinstance(results, dict):
                for key, value in results.items():
                    doc.add_heading(key, level=3)
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(str(results))
            doc.add_paragraph()  # Blank line
        
        # Chunks
        if "chunks" in data and data["chunks"]:
            doc.add_page_break()
            doc.add_heading("Source Chunks", level=2)
            for i, chunk in enumerate(data["chunks"], 1):
                doc.add_heading(f"Chunk {i}", level=3)
                chunk_text = chunk.get("text", str(chunk)) if isinstance(chunk, dict) else str(chunk)
                # Limit chunk preview to 500 chars
                if len(chunk_text) > 500:
                    chunk_text = chunk_text[:500] + "..."
                doc.add_paragraph(chunk_text)
        
        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def get_extension(self) -> str:
        return "docx"

    def get_mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
